import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import subprocess


def sanitize_filename(filename):
    """
    Poistaa tiedostonimestä kielletyt merkit ja Unicode-erikoismerkit.

    Args:
        filename (str): Alkuperäinen tiedostonimi.

    Returns:
        str: Puhdistettu tiedostonimi.
    """
    # Poistetaan kielletyt merkit ja korvataan erikoismerkit
    sanitized = re.sub(r'[<>:"/\\|?*\u0000-\u001F\u007F-\u009F\u2028\u2029]', '_', filename)
    sanitized = sanitized.replace("", "_").replace(":", "_")  # Korvaa Unicode-erikoismerkit
    sanitized = sanitized.encode('ascii', 'ignore').decode('ascii')  # Poistetaan ei-ASCII-merkit
    return sanitized[:250]



def convert_mermaid_to_image(mermaid_code, output_image):
    """
    Muuntaa Mermaid-kaavion kuvatiedostoksi käyttäen mermaid.cli-työkalua.

    Args:
        mermaid_code (str): Mermaid-kaavion koodi.
        output_image (str): Polku tallennettuun kuvatiedostoon.
    """
    with open("temp.mmd", "w", encoding="utf-8") as temp_file:
        temp_file.write(mermaid_code)

    try:
        subprocess.run(["mmdc", "-i", "temp.mmd", "-o", output_image], check=True)
    finally:
        if os.path.exists("temp.mmd"):
            os.remove("temp.mmd")


def convert_html_to_word(html_file, output_file):
    """
    Muuntaa HTML-tiedoston Word-dokumentiksi ja käsittelee Mermaid-kaaviot.

    Args:
        html_file (str): Polku HTML-tiedostoon.
        output_file (str): Polku tallennettuun Word-dokumenttiin.
    """
    try:
        with open(html_file, 'r', encoding='utf-8') as file:
            content = file.read()
    except FileNotFoundError:
        print(f"HTML-tiedostoa ei löydy: {html_file}")
        return

    soup = BeautifulSoup(content, 'html.parser')

    # Luo uusi Word-dokumentti
    document = Document()

    # Lisää HTML:n tekstit ja kaaviot dokumenttiin
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'div']):
        if element.name.startswith('h'):
            document.add_heading(element.text.strip(), level=int(element.name[1]))
        elif element.name == 'p':
            document.add_paragraph(element.text.strip())
        elif element.name == 'li':
            document.add_paragraph(f"- {element.text.strip()}")
        elif element.name == 'div' and 'mermaid' in element.get('class', []):
            mermaid_code = element.text.strip()
            image_path = "temp_mermaid.png"
            convert_mermaid_to_image(mermaid_code, image_path)
            if os.path.exists(image_path):
                document.add_picture(image_path, width=Inches(5))
                os.remove(image_path)

    # Tallenna Word-dokumentti
    document.save(output_file)


def process_directory(input_dir, output_dir):
    """
    Käy läpi kansio ja sen alikansiot ja muuntaa HTML-tiedostot Wordiksi.

    Args:
        input_dir (str): Lähtökansion polku.
        output_dir (str): Kohdekansion polku.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith('.html'):
                html_path = os.path.join(root, file)

                # Luo vastaava alikansiorakenne kohdekansioon
                relative_path = os.path.relpath(root, input_dir)
                sanitized_subdir = sanitize_filename(relative_path)
                output_subdir = os.path.join(output_dir, sanitized_subdir)

                if not os.path.exists(output_subdir):
                    os.makedirs(output_subdir)

                sanitized_filename = sanitize_filename(file)
                word_file = os.path.join(output_subdir, f"{os.path.splitext(sanitized_filename)[0]}.docx")

                try:
                    print(f"Converting {html_path} to {word_file}")
                    convert_html_to_word(html_path, word_file)

                    # Tarkistetaan, että tiedosto luotiin
                    if not os.path.exists(word_file):
                        raise FileNotFoundError(f"Tiedostoa ei luotu: {word_file}")
                except Exception as e:
                    print(f"Virhe käsiteltäessä tiedostoa {html_path}: {e}")


if __name__ == "__main__":
    root_directory = input("Anna juurikansion polku: ")
    words_directory = os.path.join(root_directory, "words")

    process_directory(root_directory, words_directory)
    print("HTML-tiedostojen muunto Wordiksi on valmis!")
