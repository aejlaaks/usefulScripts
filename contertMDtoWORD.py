import os
import markdown
import subprocess
from pathlib import Path
from markdown.extensions import extra, codehilite
from docx import Document
from docx.shared import Inches
import tempfile

def convert_markdown_to_html(markdown_content):
    """Convert Markdown content to HTML."""
    md = markdown.Markdown(extensions=[extra.ExtraExtension(), codehilite.CodeHiliteExtension()])
    return md.convert(markdown_content)

def generate_mermaid_diagram(mermaid_code, output_path):
    """Generate a Mermaid diagram as PNG using mmdc."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mmd") as temp_file:
        temp_file.write(mermaid_code.encode('utf-8'))
        temp_file.close()

        result = subprocess.run(
            ["mmdc", "-i", temp_file.name, "-o", output_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )

        if result.returncode != 0:
            print(f"Error generating Mermaid diagram: {result.stderr}")
        os.unlink(temp_file.name)  # Remove the temporary file

def add_content_to_word(markdown_content, word_path):
    """Convert Markdown content to a Word document."""
    doc = Document()

    for line in markdown_content.splitlines():
        line = line.strip()

        if line.startswith("```mermaid"):
            mermaid_code = []
            while True:
                line = next(iter(markdown_content.splitlines()), "").strip()
                if line == "```":
                    break
                mermaid_code.append(line)
            mermaid_code = "\n".join(mermaid_code)

            # Generate the Mermaid diagram
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                generate_mermaid_diagram(mermaid_code, temp_img.name)
                doc.add_picture(temp_img.name, width=Inches(5))
                os.unlink(temp_img.name)
        else:
            doc.add_paragraph(line)

    doc.save(word_path)

def generate_pdf(html_content, pdf_path):
    """Generate a PDF from HTML using Puppeteer."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as temp_html:
        temp_html.write(html_content.encode('utf-8'))
        temp_html.close()

        result = subprocess.run(
            ["node", "generate-pdf.js", temp_html.name, pdf_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )

        if result.returncode != 0:
            print(f"Error generating PDF: {result.stderr}")
        os.unlink(temp_html.name)

def process_markdown_files(input_dir):
    """Process all Markdown files in the directory."""
    pdf_output_dir = os.path.join(input_dir, "PDFs")
    word_output_dir = os.path.join(input_dir, "WordDocs")
    os.makedirs(pdf_output_dir, exist_ok=True)
    os.makedirs(word_output_dir, exist_ok=True)

    for md_file in Path(input_dir).rglob("*.md"):
        print(f"Processing file: {md_file}")

        with open(md_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Convert to HTML
        html_content = convert_markdown_to_html(markdown_content)

        pdf_path = os.path.join(pdf_output_dir, f"{md_file.stem}.pdf")
        generate_pdf(html_content, pdf_path)
        # Generate PDF

        # Generate Word document
        word_path = os.path.join(word_output_dir, f"{md_file.stem}.docx")
        add_content_to_word(markdown_content, word_path)

if __name__ == "__main__":
    input_dir = input("Anna kansion polku, jossa Markdown-tiedostot sijaitsevat: ").strip('"')
    if not os.path.exists(input_dir):
        print(f"Virhe: Hakemistoa ei löydy: {input_dir}")
    else:
        process_markdown_files(input_dir)
        print("Kaikki Markdown-tiedostot on käsitelty!")
