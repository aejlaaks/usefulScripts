import os
import re
import subprocess
from pathlib import Path
from markdown2 import markdown
from bs4 import BeautifulSoup
from docx import Document
import pdfkit

# CSS-tyylit
CSS_STYLES = """
<style>
    body { font-family: Arial, sans-serif; line-height: 1.6; color: #333; background-color: #fefefe; }
    h1 { color: #2E86C1; }
    h2 { color: #28B463; }
    h3 { color: #F39C12; }
    p { color: #555; margin: 10px 0; }
    code { background-color: #FDEBD0; col§or: #D35400; padding: 2px 4px; border-radius: 3px; }
    pre { background-color: #FDEBD0; padding: 10px; border-radius: 5px; overflow-x: auto; color: #D35400; }
    img { max-width: 100%; height: auto; border: 2px solid #ddd; border-radius: 5px; }
    ul { color: #16A085; }
    ol { color: #AF7AC5; }
    table { border-collapse: collapse; width: 100%; margin: 20px 0; background-color: #F9E79F; }
    th, td { border: 1px solid #D5DBDB; padding: 8px; text-align: left; }
    th { background-color: #FAD7A0; color: #333; }
</style>
"""

MERMAID_CLI_PATH = "C:\\Users\\Antti\\AppData\\Roaming\\npm\\mmdc.cmd"
WKHTMLTOPDF_PATH = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"


def sanitize_filename(filename):
    """Poistaa tiedostonimestä kaikki erikoismerkit ja lyhentää liian pitkät nimet."""
    sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', filename)
    return sanitized[:100] if len(sanitized) > 100 else sanitized


def render_mermaid_diagram(mermaid_code, output_dir):
    """Renderöi Mermaid-kaavion PNG-kuvaksi."""
    try:
        output_png = output_dir / f"mermaid_{hash(mermaid_code)}.png"
        temp_file = output_dir / "temp_mermaid.mmd"

        with open(temp_file, "w", encoding="utf-8") as f:
            f.write(mermaid_code)

        subprocess.run(
            [MERMAID_CLI_PATH, "-i", str(temp_file), "-o", str(output_png)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        temp_file.unlink()
        return output_png
    except subprocess.CalledProcessError as e:
        print(f"Mermaid CLI epäonnistui:\n{e}")
        return None


def process_markdown_file(file_path, html_dir, pdf_dir, word_dir, image_dir):
    """Käsittelee yksittäisen Markdown-tiedoston."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            markdown_content = f.read()

        # Käsittele Mermaid-kaaviot
        mermaid_pattern = re.compile(r"```mermaid([\s\S]*?)```")
        matches = mermaid_pattern.findall(markdown_content)

        for match in matches:
            mermaid_code = match.strip()
            rendered_image = render_mermaid_diagram(mermaid_code, image_dir)
            if rendered_image:
                markdown_content = markdown_content.replace(
                    f"```mermaid{match}```", f'<img src="{rendered_image}" alt="Mermaid Diagram"/>'
                )

        # Muunna HTML:ksi
        html_content = markdown(markdown_content)
        full_html = f"<html><head>{CSS_STYLES}</head><body>{html_content}</body></html>"

        # Puhdistettu tiedostonimi
        sanitized_name = sanitize_filename(file_path.stem)
        html_file = html_dir / f"{sanitized_name}.html"

        with open(html_file, "w", encoding="utf-8") as f:
            f.write(full_html)

        if not html_file.exists():
            print(f"HTML-tiedoston tallennus epäonnistui: {html_file}")
            return

        # PDF-muunnos
        pdf_file = pdf_dir / f"{sanitized_name}.pdf"
        config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)
        pdfkit.from_file(f"file:///{html_file.resolve()}", str(pdf_file), configuration=config)

        # Word-muunnos
        convert_html_to_word(full_html, word_dir / f"{sanitized_name}.docx")

    except Exception as e:
        print(f"Virhe tiedoston käsittelyssä ({file_path}): {e}")


def convert_html_to_word(html_content, output_path):
    """Muuntaa HTML:n Word-dokumentiksi."""
    doc = Document()
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.body.children:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "h3":
            doc.add_heading(element.text, level=3)
        elif element.name == "p":
            doc.add_paragraph(element.text)
        elif element.name == "pre":
            doc.add_paragraph(element.text, style="Code")
        elif element.name == "img":
            img_path = element["src"]
            doc.add_picture(str(img_path))
    doc.save(output_path)


def process_markdown_files(root_dir):
    """Käsittelee Markdown-tiedostot ja luo tulosteet."""
    html_dir = root_dir / "html"
    pdf_dir = root_dir / "pdf"
    word_dir = root_dir / "word"
    image_dir = root_dir / "images"

    html_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)
    word_dir.mkdir(parents=True, exist_ok=True)
    image_dir.mkdir(parents=True, exist_ok=True)

    for file in root_dir.rglob("*.md"):
        print(f"Käsitellään tiedostoa: {file}")
        process_markdown_file(file, html_dir, pdf_dir, word_dir, image_dir)


if __name__ == "__main__":
    root_path = Path(input("Anna juurikansion polku: ").strip())
    if not root_path.exists():
        print("Virheellinen kansio!")
    else:
        process_markdown_files(root_path)
